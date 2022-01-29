# %%

import docx
import os


def info_update(doc):
    '''此函数用于批量替换需要替换的信息
    doc:文件
    old_info和new_info：原文字和需要替换的新文字
    '''
    # 读取段落中的所有run，找到需替换的信息进行替换
    for para in doc.paragraphs:  #
        for run in para.runs:
            # run.text = run.text.replace(old_info, new_info)  # 替换信息-多内容替换
            run.text = replace_doc(run.text, getReplacedList())
    # 读取表格中的所有单元格，找到需替换的信息进行替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace(old_info, new_info)  # 替换信息


def read_txt(filename):
    with open(filename, "r", encoding='utf-8') as f:  # 打开文件
        data = f.read()  # 读取文件
        return data


def replace_doc(text, replacedList):
    for i in replacedList:
        if i in text:
            text = text.replace(i, replacedList[i])
    return text


# %%
"""
1.读取文件夹下所有文件
2.将所有doc文件转换成docx
3.多字段,批量修改
"""

path = "D:/data/"  # 需处理文件夹路径
files = []
for file in os.listdir(path):
    if file.endswith(".docx") or file.endswith(".doc"):  # 排除文件夹内的其它干扰文件，只获取word文件
        files.append(path + file)
print(files)
"""
替换集合处理
"""


def getReplacedList():
    replacedStr = read_txt('replace.txt').split("\n")
    replacedList = {}
    for rep in replacedStr:
        list_rep = rep.split(" ")
        replacedList[list_rep[0]] = list_rep[1]
    return replacedList


for file in files:
    if file.endswith(".docx"):
        doc = docx.Document(file)
        info_update(doc)
        doc.save("./result/{}".format(file.split("/")[-1]))  # 输出路径
        print("{}替换完成".format(file))
